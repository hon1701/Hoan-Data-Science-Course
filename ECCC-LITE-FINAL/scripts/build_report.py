"""Build the final submission report from the retained Word template."""

from __future__ import annotations

import json
import math
import shutil
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parents[1]
REFERENCE = ROOT.parent / "notes" / "Khung_bao_cao.docx"
OUTPUT = ROOT / "reports" / "BaoCao_NOP.docx"
TABLES = ROOT / "outputs" / "tables"
FIGURES = ROOT / "outputs" / "figures"

BLUE = "1F4E78"
HEADER_FILL = "D9E2F3"
LIGHT_FILL = "EEF3F8"
WHITE = "FFFFFF"
USABLE_DXA = 8_856


def vn_number(value: float, digits: int = 4) -> str:
    return f"{value:.{digits}f}".replace(".", ",")


def percent(value: float, digits: int = 2) -> str:
    return f"{value * 100:.{digits}f}%".replace(".", ",")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 100, bottom: int = 80, end: int = 100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_row_cant_split(row) -> None:
    """Keep a table row together so Word cannot split one record across pages."""
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        cant_split = OxmlElement("w:cantSplit")
        cant_split.set(qn("w:val"), "true")
        tr_pr.append(cant_split)


def set_table_geometry(table, widths: list[int]) -> None:
    if sum(widths) != USABLE_DXA:
        raise ValueError(f"Tổng độ rộng cột phải bằng {USABLE_DXA}: {widths}")
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(USABLE_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        set_row_cant_split(row)
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def format_run(run, *, bold: bool | None = None, italic: bool | None = None, size: float | None = None, color: str | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Times New Roman")
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Times New Roman")
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Times New Roman")
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_body(doc: Document, text: str, *, first_line: bool = True, italic: bool = False, bold_lead: str | None = None) -> Paragraph:
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.space_after = Pt(4)
    if first_line:
        paragraph.paragraph_format.first_line_indent = Inches(0.5)
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        format_run(lead, bold=True, italic=italic, size=13)
        rest = paragraph.add_run(text[len(bold_lead):])
        format_run(rest, italic=italic, size=13)
    else:
        run = paragraph.add_run(text)
        format_run(run, italic=italic, size=13)
    return paragraph


def add_bullet(doc: Document, text: str) -> Paragraph:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(text)
    format_run(run, size=13)
    return paragraph


def add_number(doc: Document, text: str) -> Paragraph:
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(text)
    format_run(run, size=13)
    return paragraph


def add_heading(doc: Document, text: str, level: int = 1, *, page_break: bool = False) -> Paragraph:
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    if page_break:
        paragraph.paragraph_format.page_break_before = True
    paragraph.paragraph_format.keep_with_next = True
    for run in paragraph.runs:
        format_run(run, bold=True, size=15 if level == 1 else 13.5, color="000000")
    return paragraph


def add_caption(doc: Document, text: str) -> Paragraph:
    paragraph = doc.add_paragraph(style="Caption")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    format_run(run, bold=True, size=10.5, color=BLUE)
    return paragraph


def add_figure(doc: Document, filename: str, caption: str, interpretation: str, *, width: float = 5.9) -> None:
    path = FIGURES / filename
    if not path.is_file():
        raise FileNotFoundError(path)
    add_caption(doc, caption)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    shape = paragraph.add_run().add_picture(str(path), width=Inches(width))
    shape._inline.docPr.set("descr", caption)
    note = add_body(doc, interpretation, first_line=False, italic=True)
    note.paragraph_format.space_after = Pt(7)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int], *, font_size: float = 10.5) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = 0
    header = table.rows[0]
    set_repeat_table_header(header)
    for cell, text in zip(header.cells, headers):
        cell.text = text
        set_cell_shading(cell, HEADER_FILL)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                format_run(run, bold=True, size=font_size)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for column, (cell, text) in enumerate(zip(cells, values)):
            cell.text = str(text)
            if row_index % 2:
                set_cell_shading(cell, LIGHT_FILL)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if column > 0 else WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    format_run(run, size=font_size)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_toc_before(section_break, doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("MỤC LỤC")
    format_run(run, bold=True, size=16)
    section_break.addprevious(title._p)

    paragraph = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), 'TOC \\o "1-2" \\h \\z \\u')
    run_node = OxmlElement("w:r")
    text_node = OxmlElement("w:t")
    text_node.text = "Mục lục sẽ được cập nhật tự động khi mở bằng Microsoft Word."
    run_node.append(text_node)
    fld.append(run_node)
    paragraph._p.append(fld)
    section_break.addprevious(paragraph._p)


def set_update_fields(doc: Document) -> None:
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def clear_submission_body(doc: Document):
    """Keep cover/review and first section break, remove draft report body."""

    body = doc._element.body
    children = list(body.iterchildren())
    section_break = None
    for child in children:
        if child.tag == qn("w:p") and child.find(".//" + qn("w:sectPr")) is not None:
            section_break = child
            break
    if section_break is None:
        raise AssertionError("Không tìm thấy section break của template.")
    break_index = children.index(section_break)
    final_sect_pr = children[-1]
    for index, child in enumerate(children):
        if child is final_sect_pr or child is section_break:
            continue
        if index >= 23:
            body.remove(child)
    return section_break


def update_cover_date(doc: Document) -> None:
    cover = doc.tables[0]
    for paragraph in cover.cell(0, 0).paragraphs:
        if "TP.HCM" in paragraph.text:
            paragraph.clear()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run("TP.HCM, Tháng 09 Năm 2026")
            format_run(run, bold=True, italic=True, size=13)


def build() -> Path:
    if not REFERENCE.is_file():
        raise FileNotFoundError(REFERENCE)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(REFERENCE, OUTPUT)
    doc = Document(OUTPUT)
    update_cover_date(doc)
    section_break = clear_submission_body(doc)
    add_toc_before(section_break, doc)
    set_update_fields(doc)

    audit = pd.read_csv(TABLES / "data_audit.csv").set_index("metric")["value"]
    split = pd.read_csv(TABLES / "split_summary.csv")
    train_summary = pd.read_csv(TABLES / "train_class_summary.csv")
    correlations = pd.read_csv(TABLES / "selected_correlations.csv")
    model_candidates = pd.read_csv(TABLES / "model_candidates.csv")
    comparison = pd.read_csv(TABLES / "model_comparison.csv")
    top_p = pd.read_csv(TABLES / "top_p_metrics.csv")
    errors = pd.read_csv(TABLES / "error_examples.csv")
    importance = pd.read_csv(TABLES / "feature_importance.csv")
    modeling = json.loads((TABLES / "modeling_summary.json").read_text(encoding="utf-8"))
    evaluation = json.loads((TABLES / "evaluation_summary.json").read_text(encoding="utf-8"))

    val = evaluation["validation"]
    test = evaluation["test"]
    primary = test["cost_threshold_metrics"]
    rf_val = comparison[(comparison["split"] == "validation") & (comparison["model"] == "random_forest")].iloc[0]
    log_val = comparison[(comparison["split"] == "validation") & (comparison["model"] == "logistic")].iloc[0]
    dummy_val = comparison[(comparison["split"] == "validation") & (comparison["model"] == "dummy")].iloc[0]

    # PHẦN MỞ ĐẦU
    add_heading(doc, "PHẦN MỞ ĐẦU", 1)
    add_heading(doc, "1. Lý do chọn đề tài", 2)
    add_body(doc, "Gian lận thẻ là bài toán có lớp dương rất hiếm trong khi nguồn lực kiểm tra có giới hạn. Nếu chỉ tối ưu Accuracy, một mô hình dự đoán mọi giao dịch là hợp lệ vẫn có thể đạt kết quả bề ngoài rất cao nhưng không phát hiện được gian lận. Vì vậy, đề tài tập trung vào khả năng xếp hạng giao dịch đáng ngờ và lượng fraud thu hồi được trong một ngân sách kiểm tra xác định.")
    add_body(doc, f"Bộ dữ liệu gốc có 492 giao dịch gian lận trên 284.807 dòng, tương đương {percent(492/284807, 3)}. Mức mất cân bằng này đòi hỏi Average Precision (AP), Precision, Recall, F1 và Top-p thay vì chỉ một con số Accuracy.")

    add_heading(doc, "2. Mục tiêu và phát biểu bài toán", 2)
    for item in [
        "Kiểm chứng nguồn, schema, missing, infinity, dòng trùng và phân bố Class của dữ liệu.",
        "Thực hiện EDA trên train, tạo LogAmount và giữ source_row để truy vết.",
        "So sánh DummyClassifier, Logistic Regression và Random Forest trên cùng split.",
        "Chọn mô hình bằng AP validation; chọn threshold trên validation trước khi mở test.",
        "Báo cáo AP, ROC-AUC, Precision, Recall, F1, confusion matrix và Top-0,5%/1%/2%.",
        "Tạo quy trình chạy lại được và mọi số liệu truy về artifact trong outputs/tables.",
    ]:
        add_bullet(doc, item)
    add_caption(doc, "Bảng 0.1. Phát biểu bài toán theo ba tầng")
    add_table(
        doc,
        ["Tầng", "Phát biểu"],
        [
            ["Business Problem", "Ưu tiên giao dịch đáng ngờ để giảm bỏ sót trong nguồn lực kiểm tra có hạn."],
            ["Data Science Question", "Từ V1-V28, Time và Amount, có thể tạo score xếp hạng nguy cơ gian lận hay không?"],
            ["Machine Learning Problem", "Phân loại nhị phân mất cân bằng; score của Class=1 dùng cho đánh giá và xếp hạng."],
        ],
        [2_100, 6_756],
    )

    add_heading(doc, "3. Đối tượng, phạm vi và phương pháp nghiên cứu", 2)
    add_body(doc, "Đối tượng nghiên cứu là các giao dịch trong bộ Credit Card Fraud Detection do MLG-ULB công bố. V1-V28 là các thành phần đã được biến đổi và ẩn danh; Time là số giây tương đối từ giao dịch đầu tiên; Amount là số tiền; Class=1 biểu thị gian lận.")
    add_body(doc, "Phạm vi là bài kết thúc học phần, không xây API/website và không xem kết quả là mô hình sẵn sàng triển khai trong ngân hàng. Nhóm kiểm tra hash, làm sạch, chia train/validation/test có phân tầng, EDA trên train, huấn luyện mô hình, khóa quyết định trên validation và chỉ đánh giá test một lần.")

    # CHƯƠNG 1
    add_heading(doc, "CHƯƠNG 1: CƠ SỞ LÝ THUYẾT VÀ DỮ LIỆU", 1, page_break=True)
    add_heading(doc, "1.1. Bài toán xếp hạng rủi ro và đặc điểm mất cân bằng", 2)
    add_body(doc, "Mô hình tạo score cho lớp gian lận. Score cao hơn nghĩa là giao dịch được xếp đáng ngờ hơn; score trong đề tài chủ yếu dùng để xếp hạng tương đối, không được diễn giải mặc định là xác suất đã hiệu chỉnh. Hai cách sử dụng là cắt theo threshold để tạo cảnh báo và sắp giảm dần để chọn Top-p.")
    add_body(doc, "False Negative là fraud bị bỏ sót; False Positive là giao dịch hợp lệ bị cảnh báo. Hai loại lỗi có hậu quả khác nhau, vì vậy kết quả luôn báo cả Precision và Recall cùng số TP, FP, FN, TN.")

    add_heading(doc, "1.2. Cơ sở lý thuyết các mô hình và chỉ số", 2)
    add_body(doc, "Logistic Regression học tổ hợp tuyến tính z = β₀ + Σβⱼxⱼ và ánh xạ qua sigmoid σ(z)=1/(1+e⁻ᶻ). Mô hình được đặt trong pipeline với StandardScaler để mọi split đi qua đúng cùng phép biến đổi học từ train.")
    add_body(doc, "Random Forest kết hợp nhiều cây quyết định huấn luyện trên mẫu bootstrap và tập con feature ngẫu nhiên. Score Class=1 là trung bình score của các cây. Mô hình có thể biểu diễn quan hệ phi tuyến nhưng feature importance chỉ thể hiện mức phụ thuộc của mô hình, không chứng minh quan hệ nhân quả.")
    add_body(doc, "Average Precision được tính theo AP = Σₙ(Rₙ - Rₙ₋₁)Pₙ. AP đánh giá chất lượng xếp hạng trên toàn dải threshold và phù hợp hơn ROC-AUC khi lớp dương rất hiếm; no-skill baseline của mỗi split bằng tỷ lệ Class=1 trong chính split đó.")

    add_heading(doc, "1.3. Dữ liệu và mô hình sử dụng", 2)
    add_caption(doc, "Bảng 1.1. Nhóm biến và giới hạn diễn giải")
    add_table(
        doc,
        ["Nhóm biến", "Ý nghĩa", "Giới hạn diễn giải"],
        [
            ["V1-V28", "Các thành phần số đã PCA/ẩn danh", "Không gán thành tuổi, vị trí hay hành vi cụ thể"],
            ["Time", "Số giây tương đối", "Không phải ngày giờ thật; chỉ phủ khoảng 48 giờ"],
            ["Amount", "Số tiền giao dịch", "Lệch phải; tạo LogAmount=log1p(Amount)"],
            ["Class", "0 hợp lệ; 1 gian lận", "Target, tuyệt đối không đưa vào feature"],
            ["source_row", "Khóa truy vết về raw", "Không phải feature"],
        ],
        [1_600, 3_100, 4_156],
    )
    add_caption(doc, "Bảng 1.2. Vai trò ba họ mô hình")
    add_table(
        doc,
        ["Mô hình", "Vai trò", "Thiết lập chính"],
        [
            ["DummyClassifier", "Baseline", "most_frequent minh họa Accuracy; stratified tạo baseline score"],
            ["Logistic Regression", "Mô hình tuyến tính", "StandardScaler, C=1; thử standard/balanced"],
            ["Random Forest", "Mô hình phi tuyến", "160 cây, min_samples_leaf=2; thử standard/balanced"],
        ],
        [2_000, 2_200, 4_656],
    )

    add_heading(doc, "1.4. Chỉ số đánh giá, rủi ro và giới hạn", 2)
    add_caption(doc, "Bảng 1.3. Chỉ số và câu hỏi trả lời")
    add_table(
        doc,
        ["Chỉ số", "Câu hỏi trả lời"],
        [
            ["AP", "Mô hình xếp hạng fraud tốt đến mức nào trên toàn dải threshold?"],
            ["ROC-AUC", "Khả năng phân biệt hai lớp theo cặp dương/âm?"],
            ["Precision", "Trong các cảnh báo, tỷ lệ fraud thật là bao nhiêu?"],
            ["Recall", "Trong toàn bộ fraud, mô hình phát hiện được bao nhiêu?"],
            ["F1", "Mức cân bằng Precision-Recall tại threshold cụ thể?"],
            ["Top-p và Lift", "Với ngân sách kiểm tra p%, thu hồi bao nhiêu fraud và tập trung hơn ngẫu nhiên bao nhiêu lần?"],
        ],
        [1_700, 7_156],
    )
    for risk in [
        "Data leakage: Class/source_row hoặc bước fit từ validation/test lọt vào feature pipeline.",
        "Selection bias: thay model hoặc threshold sau khi xem test.",
        "Giới hạn dữ liệu: V1-V28 ẩn danh, thiếu lịch sử khách hàng và chỉ khoảng hai ngày.",
        "Giới hạn triển khai: giả định chi phí mang tính học thuật, chưa phải quy định ngân hàng.",
    ]:
        add_bullet(doc, risk)

    # CHƯƠNG 2
    add_heading(doc, "CHƯƠNG 2: PHƯƠNG PHÁP VÀ QUY TRÌNH XÂY DỰNG", 1, page_break=True)
    add_heading(doc, "2.1. Luồng thực hiện, nguồn và kiểm chứng dữ liệu", 2)
    add_body(doc, "Quy trình gồm 01_data_eda.ipynb, 02_modeling.ipynb, 03_evaluation.ipynb và Fraud_Project_Final.ipynb. Mỗi notebook ghi artifact ổn định vào outputs/tables, outputs/figures hoặc outputs/models để kết quả trong báo cáo có thể truy ngược.")
    add_caption(doc, "Bảng 2.1. Manifest dữ liệu nguồn")
    add_table(
        doc,
        ["Kiểm tra", "Giá trị đã xác minh"],
        [
            ["Tên file", "creditcard.csv"],
            ["Kích thước", "150.828.752 byte"],
            ["SHA-256", "76274b691b16a6c49d3f159c883398e03ccd6d1ee12d9d8ee38f4b4b98551a89"],
            ["Shape", "284.807 dòng x 31 cột"],
            ["Phân bố raw", "Class 0 = 284.315; Class 1 = 492"],
        ],
        [2_000, 6_856],
        font_size=9.5,
    )

    add_heading(doc, "2.2. Làm sạch và chia dữ liệu", 2)
    add_body(doc, "source_row được tạo trước khi lọc để giữ truy vết. Nhóm kiểm tra missing, infinity và exact duplicates trên 31 cột gốc; loại 1.081 dòng trùng hoàn toàn, trong đó có 19 fraud. Đây là quy ước học thuật nhằm tránh cùng một mẫu xuất hiện ở nhiều split, không khẳng định mọi dòng trùng là lỗi nghiệp vụ.")
    split_rows = [["Raw", "284.807", "284.315", "492"], ["Sau bỏ trùng", "283.726", "283.253", "473"]]
    for _, row in split.iterrows():
        label = {"train": "Train 60%", "validation": "Validation 20%", "test": "Test 20%"}[row["split"]]
        split_rows.append([label, f"{int(row['rows']):,}".replace(",", "."), f"{int(row['class_0']):,}".replace(",", "."), str(int(row["class_1"]))])
    add_caption(doc, "Bảng 2.2. Phân bố dữ liệu sau làm sạch và chia tập")
    add_table(doc, ["Tập", "Tổng dòng", "Class 0", "Class 1"], split_rows, [2_300, 2_100, 2_100, 2_356])
    add_body(doc, "Hai lần train_test_split đều stratify=Class và random_state=42: tách test 20% trước, sau đó tách validation 25% từ phần 80% còn lại. Ba tập không giao nhau theo source_row; test được khóa tới bước đánh giá cuối.")

    add_heading(doc, "2.3. EDA và câu hỏi phân tích", 2)
    for item in [
        "Mức mất cân bằng của Class và hệ quả với Accuracy.",
        "Phân phối Amount/LogAmount theo Class; không tự động coi outlier là lỗi.",
        "Phân bố Time theo Class với giới hạn Time không phải ngày giờ thật.",
        "Tương quan Pearson chọn lọc trên train; không suy ra quan hệ nhân quả.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "2.4. Chuẩn bị feature và chống leakage", 2)
    add_body(doc, "Feature cốt lõi gồm Time, V1-V28 và LogAmount=log1p(Amount). Amount gốc được thay bằng LogAmount; Class và source_row bị loại. Mã dừng nếu tập feature thiếu hoặc dư cột. StandardScaler của Logistic Regression chỉ fit trên train; Random Forest không scale.")

    add_heading(doc, "2.5. Mô hình và quy tắc xuất score", 2)
    add_body(doc, "Dummy most_frequent minh họa Accuracy cao nhưng Recall fraud bằng 0. Dummy stratified dùng random_state=42 để tạo baseline score. Logistic Regression và Random Forest đều xuất predict_proba(X)[:,1] sau khi xác nhận classes_ chứa đúng lớp 1.")
    add_body(doc, "Trước khi ghi validation_scores.csv, score được kiểm tra đúng số dòng, hữu hạn, nằm trong [0,1]; score Logistic/Random Forest phải có hơn 10 giá trị khác nhau để phát hiện trường hợp xuất nhầm nhãn 0/1.")
    add_caption(doc, "Bảng 2.3. Candidate được thử có kiểm soát")
    candidate_rows = []
    for _, row in model_candidates.iterrows():
        candidate_rows.append([row["family"], row["candidate"], vn_number(float(row["validation_ap"])), vn_number(float(row["validation_roc_auc"]))])
    add_table(doc, ["Họ mô hình", "Candidate", "AP validation", "ROC-AUC"], candidate_rows, [2_400, 2_200, 2_128, 2_128])

    add_heading(doc, "2.6. Chọn mô hình, threshold và test cuối", 2)
    for item in [
        "Chọn candidate tốt nhất của từng họ bằng AP validation.",
        "So sánh Logistic Regression và Random Forest; nếu chênh AP dưới 0,01, chọn Logistic Regression theo tie-break đã chốt.",
        "Trên score validation của model đã chọn, quét threshold và tối thiểu hóa ExpectedCost(t)=[20*FN(t)+1*FP(t)]/Nval; nếu hòa, ưu tiên Recall cao hơn.",
        "Tính thêm threshold tối đa F1 để đối chiếu; trong lần chạy này hai quy tắc cho cùng threshold.",
        "Khóa feature, model, hyperparameter và threshold trước khi mở test; test chỉ đánh giá một lần.",
        "Ước lượng khoảng tin cậy 95% của AP test bằng 1.000 bootstrap sample, random_state=42.",
    ]:
        add_number(doc, item)

    add_heading(doc, "2.7. Top-p, Lift và phân tích FP/FN", 2)
    add_body(doc, "Với N giao dịch, k=ceil(pN). Giao dịch được sắp theo score giảm dần; source_row tăng dần là tie-break. Precision@k=TP/k; Recall@k=TP/N₁; Lift@k=Precision@k/(N₁/N). Nhóm dùng p=0,5%, 1% và 2%.")
    add_body(doc, "Sau đánh giá định lượng, nhóm xem 10 False Positive có score cao nhất và 10 False Negative có score thấp nhất. Nhận xét chỉ mô tả mẫu số học và không gán nguyên nhân nghiệp vụ cho biến ẩn danh.")

    # CHƯƠNG 3
    add_heading(doc, "CHƯƠNG 3: THỰC NGHIỆM VÀ ĐÁNH GIÁ KẾT QUẢ", 1, page_break=True)
    add_heading(doc, "3.1. Thiết kế thực nghiệm và kiểm tra dữ liệu", 2)
    env = evaluation["environment"]
    add_caption(doc, "Bảng 3.1. Môi trường chạy cuối")
    add_table(
        doc,
        ["Thành phần", "Giá trị"],
        [
            ["Python", env["python"]],
            ["NumPy / pandas", f"{env['numpy']} / {env['pandas']}"],
            ["scikit-learn", env["scikit_learn"]],
            ["matplotlib", env["matplotlib"]],
            ["Seed", "42"],
            ["Dữ liệu", "creditcard.csv đã xác minh SHA-256"],
        ],
        [2_300, 6_556],
    )
    add_caption(doc, "Bảng 3.2. Kết quả audit lần chạy cuối")
    add_table(
        doc,
        ["Kiểm tra", "Kết quả"],
        [
            ["Raw", "284.807 dòng; 31 cột; 492 fraud"],
            ["Missing / infinity", "0 / 0"],
            ["Dòng trùng bị loại", "1.081 dòng; gồm 19 fraud"],
            ["Sau bỏ trùng", "283.726 dòng; 473 fraud"],
            ["Train / validation / test", "170.235 / 56.745 / 56.746"],
            ["Fraud theo split", "284 / 94 / 95"],
            ["Raw SHA-256", str(audit["source_sha256"])],
        ],
        [2_600, 6_256],
        font_size=9.5,
    )

    add_heading(doc, "3.2. Kết quả EDA", 2)
    train_rate = int(split.loc[split["split"] == "train", "class_1"].iloc[0]) / int(split.loc[split["split"] == "train", "rows"].iloc[0])
    add_figure(doc, "class_distribution.png", "Hình 3.1. Phân bố Class trên train", f"Train chỉ có 284 fraud trên 170.235 giao dịch ({percent(train_rate, 3)}), xác nhận Accuracy đơn lẻ không phù hợp để chọn mô hình.")
    class0 = train_summary[train_summary["Class"] == 0].iloc[0]
    class1 = train_summary[train_summary["Class"] == 1].iloc[0]
    add_figure(doc, "amount_by_class.png", "Hình 3.2. Amount/LogAmount theo Class", f"Amount trung vị của Class 0 là {vn_number(float(class0['amount_median']),2)}, cao hơn Class 1 ({vn_number(float(class1['amount_median']),2)}), trong khi mean của Class 1 cao hơn do đuôi phải. LogAmount giúp nén miền giá trị; kết quả không chứng minh Amount gây ra gian lận.")
    add_figure(doc, "time_by_class.png", "Hình 3.3. Time theo Class", f"Fraud trong train xuất hiện từ Time={int(class1['time_min']):,} tới {int(class1['time_max']):,} giây. Time chỉ là thời gian tương đối trong khoảng quan sát, không phải ngày giờ giao dịch thật.".replace(",", "."))
    top_corr = correlations.iloc[0]
    add_figure(doc, "selected_correlations.png", "Hình 3.4. Tương quan chọn lọc trên train", f"V17 có |tương quan Pearson| lớn nhất trong bảng chọn lọc ({vn_number(float(top_corr['correlation_with_class']))}). Đây chỉ là liên hệ thống kê trên train; V1-V28 không được gán ý nghĩa nghiệp vụ.")

    add_heading(doc, "3.3. Kết quả trên validation và lựa chọn mô hình", 2)
    add_caption(doc, "Bảng 3.3. So sánh validation và metric tại threshold tối đa F1")
    add_table(
        doc,
        ["Model", "AP", "ROC-AUC", "AP/Baseline", "Threshold", "Precision", "Recall", "F1"],
        [
            ["Dummy stratified", vn_number(float(dummy_val.ap)), vn_number(float(dummy_val.roc_auc)), vn_number(float(dummy_val.ap_over_baseline),1)+"x", vn_number(float(dummy_val.threshold),4), vn_number(float(dummy_val.precision)), vn_number(float(dummy_val.recall)), vn_number(float(dummy_val.f1))],
            ["Logistic", vn_number(float(log_val.ap)), vn_number(float(log_val.roc_auc)), vn_number(float(log_val.ap_over_baseline),1)+"x", vn_number(float(log_val.threshold),4), vn_number(float(log_val.precision)), vn_number(float(log_val.recall)), vn_number(float(log_val.f1))],
            ["Random Forest", vn_number(float(rf_val.ap)), vn_number(float(rf_val.roc_auc)), vn_number(float(rf_val.ap_over_baseline),1)+"x", vn_number(float(rf_val.threshold),4), vn_number(float(rf_val.precision)), vn_number(float(rf_val.recall)), vn_number(float(rf_val.f1))],
        ],
        [1_750, 950, 1_050, 1_200, 1_150, 950, 900, 906],
        font_size=8.5,
    )
    add_body(doc, f"Random Forest đạt AP validation {vn_number(float(rf_val.ap))}, cao hơn Logistic Regression {vn_number(float(log_val.ap))} một khoảng {vn_number(abs(float(rf_val.ap)-float(log_val.ap)))}. Chênh lệch vượt 0,01 nên nhóm chọn Random Forest theo AP, không cần kích hoạt tie-break. Baseline validation chỉ {vn_number(val['baseline_ap'],6)}.")
    add_figure(doc, "validation_pr_curve.png", "Hình 3.5. Precision-Recall trên validation", "Đường PR cho thấy Random Forest duy trì Precision cao hơn trong phần lớn miền Recall có ý nghĩa, phù hợp với AP validation cao nhất.")

    add_heading(doc, "3.4. Kết quả test cuối", 2, page_break=True)
    add_caption(doc, "Bảng 3.4a. Hiệu quả xếp hạng của model đã khóa trên test")
    add_table(
        doc,
        ["Model", "AP (95% CI)", "ROC-AUC", "Baseline", "AP/Baseline"],
        [[
            "Random Forest",
            f"{vn_number(test['average_precision'])} ({vn_number(test['ap_bootstrap_95_ci'][0])}-{vn_number(test['ap_bootstrap_95_ci'][1])})",
            vn_number(test["roc_auc"]),
            vn_number(test["baseline_ap"],6),
            vn_number(test["ap_over_baseline"],1)+"x",
        ]],
        [1_700, 2_456, 1_600, 1_500, 1_600],
        font_size=9,
    )
    add_caption(doc, "Bảng 3.4b. Kết quả tại threshold đã khóa")
    add_table(
        doc,
        ["Threshold", "Precision", "Recall", "F1", "TP", "FP", "FN", "TN"],
        [[
            vn_number(primary["threshold"],6),
            vn_number(primary["precision"]),
            vn_number(primary["recall"]),
            vn_number(primary["f1"]),
            str(primary["tp"]),
            str(primary["fp"]),
            str(primary["fn"]),
            str(primary["tn"]),
        ]],
        [1_656, 1_200, 1_100, 1_000, 950, 950, 950, 1_050],
        font_size=9,
    )
    add_body(doc, f"Trên test đã khóa, AP={vn_number(test['average_precision'])}, cao gấp {vn_number(test['ap_over_baseline'],1)} lần no-skill baseline. Khoảng tin cậy bootstrap 95% là [{vn_number(test['ap_bootstrap_95_ci'][0])}; {vn_number(test['ap_bootstrap_95_ci'][1])}]. Tại threshold {vn_number(primary['threshold'],6)}, mô hình phát hiện {primary['tp']}/95 fraud và tạo {primary['fp']} cảnh báo nhầm.")
    add_figure(doc, "test_pr_curve.png", "Hình 3.6. Precision-Recall trên test đã khóa", "AP được tính trực tiếp từ score liên tục bằng average_precision_score; test không được dùng để đổi model hoặc threshold.")
    add_figure(doc, "test_confusion_matrix.png", "Hình 3.7. Confusion matrix tại threshold đã khóa", f"Mô hình có TP={primary['tp']}, FP={primary['fp']}, FN={primary['fn']}, TN={primary['tn']}. Precision cao nhưng vẫn bỏ sót {primary['fn']} fraud, cho thấy threshold luôn là một đánh đổi vận hành.")

    add_heading(doc, "3.5. Kết quả Top-p", 2)
    top_rows = []
    for _, row in top_p.iterrows():
        top_rows.append([
            f"{vn_number(float(row['top_p_percent']),1)}%",
            f"{int(row['k']):,}".replace(",", "."),
            str(int(row["tp"])),
            str(int(row["fp"])),
            percent(float(row["precision_at_k"]),2),
            percent(float(row["recall_at_k"]),2),
            vn_number(float(row["lift_at_k"]),2)+"x",
        ])
    add_caption(doc, "Bảng 3.5. Hiệu quả theo năng lực kiểm tra")
    add_table(doc, ["Top-p", "k", "TP", "FP", "Precision@k", "Recall@k", "Lift@k"], top_rows, [1_050, 1_050, 800, 900, 1_700, 1_700, 1_656], font_size=9)
    add_body(doc, f"Top-0,5% kiểm tra 284 giao dịch và thu hồi 76/95 fraud ({percent(float(top_p.iloc[0]['recall_at_k']),2)}), Lift={vn_number(float(top_p.iloc[0]['lift_at_k']),2)}x. Tăng lên Top-2% chỉ thu thêm 5 fraud nhưng tạo thêm 846 FP so với Top-0,5%; lựa chọn p vì vậy phụ thuộc năng lực xử lý cảnh báo.")
    add_figure(doc, "top_p_performance.png", "Hình 3.8. Recall và Lift theo Top-p", "Recall tăng chậm khi mở rộng p, trong khi Lift giảm mạnh vì danh sách chứa nhiều giao dịch hợp lệ hơn.")

    add_heading(doc, "3.6. Phân tích False Positive, False Negative và feature importance", 2)
    fp = errors[errors["error_type"] == "False Positive"]
    fn = errors[errors["error_type"] == "False Negative"]
    add_caption(doc, "Bảng 3.6. Mẫu lỗi được kiểm tra")
    add_table(
        doc,
        ["Nhóm lỗi", "Số mẫu", "Dấu hiệu số học trong mẫu", "Giới hạn kết luận"],
        [
            ["False Positive", str(len(fp)), f"Score {vn_number(fp['score'].min(),3)}-{vn_number(fp['score'].max(),3)}; Amount {vn_number(fp['Amount'].min(),2)}-{vn_number(fp['Amount'].max(),2)}", "Không biết ngữ cảnh giao dịch/khách hàng"],
            ["False Negative", str(len(fn)), f"10 FN score thấp nhất đều bằng {vn_number(fn['score'].max(),1)}; Amount {vn_number(fn['Amount'].min(),2)}-{vn_number(fn['Amount'].max(),2)}", "Không suy ra nguyên nhân gian lận"],
        ],
        [1_500, 1_000, 3_700, 2_656],
        font_size=9.5,
    )
    top3 = importance.head(3)
    add_body(doc, f"Random Forest phụ thuộc nhiều nhất vào {top3.iloc[0]['feature']} ({vn_number(top3.iloc[0]['importance'])}), {top3.iloc[1]['feature']} ({vn_number(top3.iloc[1]['importance'])}) và {top3.iloc[2]['feature']} ({vn_number(top3.iloc[2]['importance'])}). Các giá trị này không cho phép gán ý nghĩa nghiệp vụ hoặc kết luận nhân quả.")
    add_figure(doc, "feature_importance.png", "Hình 3.9. Feature importance của Random Forest", "Thứ hạng importance nhất quán với các biến có tương quan train nổi bật, nhưng được diễn giải như mức mô hình sử dụng biến, không phải nguyên nhân fraud.")

    add_heading(doc, "3.7. Thảo luận kết quả và giới hạn", 2)
    add_body(doc, f"Mô hình tốt hơn rõ rệt so với no-skill baseline: AP test {vn_number(test['average_precision'])} so với {vn_number(test['baseline_ap'],6)}. Ở threshold chi phí, Precision {percent(primary['precision'],2)} và Recall {percent(primary['recall'],2)} tạo một danh sách cảnh báo nhỏ (83 giao dịch) nhưng bỏ sót 22 fraud. Top-0,5% thu hồi thêm ba fraud so với threshold nhưng cần kiểm tra 284 giao dịch.")
    add_body(doc, "Kết quả cho thấy score có ích để ưu tiên giao dịch trong bộ dữ liệu này, nhưng không chứng minh khả năng triển khai thực tế. Dữ liệu chỉ khoảng hai ngày, biến chính đã ẩn danh, không có quan hệ theo khách hàng/thẻ, chưa đánh giá drift và chi phí FN/FP chỉ là giả định học thuật.")

    # KẾT LUẬN
    add_heading(doc, "PHẦN KẾT LUẬN", 1, page_break=True)
    add_body(doc, "Đề tài đã hoàn thiện quy trình từ kiểm chứng nguồn, làm sạch, chia dữ liệu, EDA, xây feature, huấn luyện, lựa chọn validation đến đánh giá test và Top-p. Các bước fit không sử dụng validation/test ngoài vai trò đã định; mọi số trong báo cáo truy được về output của cùng lần chạy cuối.")
    add_body(doc, f"Nhóm chọn Random Forest vì AP validation {vn_number(float(rf_val.ap))} cao hơn Logistic Regression {vn_number(float(log_val.ap))}. Trên test, Random Forest đạt AP={vn_number(test['average_precision'])}, ROC-AUC={vn_number(test['roc_auc'])}; tại threshold {vn_number(primary['threshold'],6)}, Precision={percent(primary['precision'],2)}, Recall={percent(primary['recall'],2)} và F1={vn_number(primary['f1'])}.")
    top1 = top_p.iloc[1]
    add_body(doc, f"Với năng lực kiểm tra Top-1%, mô hình đưa 568 giao dịch lên đầu và tìm được {int(top1['tp'])}/95 fraud, Recall@Top-1%={percent(float(top1['recall_at_k']),2)}, Lift={vn_number(float(top1['lift_at_k']),2)}x. Đầu ra phù hợp để hỗ trợ xếp hạng ưu tiên, không tự động khóa thẻ.")
    add_body(doc, "Giá trị chính của đề tài là quy trình tái lập được, metric phù hợp với lớp hiếm và cách liên hệ kết quả với năng lực kiểm tra. Bước tiếp theo nên dùng dữ liệu theo thời gian dài hơn, thêm ngữ cảnh khách hàng/giao dịch, đánh giá drift và hiệu chỉnh chi phí cùng chuyên gia nghiệp vụ.")

    # PHỤ LỤC A
    appendix_a_heading = add_heading(doc, "PHỤ LỤC A: TỔ CHỨC REPOSITORY VÀ NOTEBOOK", 1, page_break=True)
    for run in appendix_a_heading.runs:
        format_run(run, bold=True, size=13)
    add_caption(doc, "Bảng A.1. File chính và vai trò")
    add_table(
        doc,
        ["Đường dẫn", "Vai trò"],
        [
            ["README.md", "Nguồn, quy tắc khoa học, cách chạy và artifact"],
            ["notebooks/01_data_eda.ipynb", "Audit, làm sạch, split, EDA"],
            ["notebooks/02_modeling.ipynb", "Feature, candidate, pipeline, validation score"],
            ["notebooks/03_evaluation.ipynb", "Threshold, test, Top-p, FP/FN"],
            ["notebooks/Fraud_Project_Final.ipynb", "Tổng hợp artifact của lần chạy cuối"],
            ["src/modeling.py", "Huấn luyện và lựa chọn model trên validation"],
            ["src/evaluation.py", "Threshold, bootstrap, test và hình đánh giá"],
            ["scripts/verify_project.py", "Kiểm tra contract toàn dự án"],
        ],
        [3_700, 5_156],
        font_size=9.5,
    )
    for heading, points in [
        ("A.1. 01_data_eda.ipynb - Hoan", ["Xác minh SHA-256/schema.", "Tạo source_row, bỏ trùng, split 60/20/20.", "EDA chỉ trên train và xuất bốn hình."]),
        ("A.2. 02_modeling.ipynb - Huy", ["Đọc đúng split A.1, không chia lại.", "Thử Dummy/Logistic/Random Forest có kiểm soát.", "Xuất validation score và khóa model trước test."]),
        ("A.3. 03_evaluation.ipynb - Sang", ["Chọn threshold trên validation.", "Đánh giá test một lần, bootstrap AP.", "Tính Top-p, FP/FN và feature importance."]),
        ("A.4. Fraud_Project_Final.ipynb - cả nhóm", ["Tổng hợp audit, model, test và hình.", "Kiểm tra artifact của cùng lần chạy.", "Nêu giới hạn và quyết định sử dụng score."]),
    ]:
        add_heading(doc, heading, 2)
        for point in points:
            add_bullet(doc, point)

    # PHỤ LỤC B
    add_heading(doc, "PHỤ LỤC B: PHÂN CÔNG VÀ TRẠNG THÁI HOÀN THÀNH", 1, page_break=True)
    add_table(
        doc,
        ["Thành viên", "Phần phụ trách", "Trạng thái"],
        [
            ["Phạm Gia Huy", "Mở đầu, cơ sở lý thuyết, 02_modeling", "Hoàn tất"],
            ["Nguyễn Văn Hoan", "Dữ liệu/EDA, phương pháp, ghép repository", "Hoàn tất"],
            ["Bùi Thái Sang", "03_evaluation, kết quả, kết luận", "Hoàn tất"],
            ["Cả nhóm", "Fraud_Project_Final, rà soát và báo cáo", "Hoàn tất"],
        ],
        [2_100, 5_156, 1_600],
    )
    add_body(doc, "Bộ bàn giao đã qua 8 unit test và verifier toàn dự án: đúng hash/split, không giao nhau, feature không chứa cột cấm, score đúng miền, model/threshold khóa trước test, Top-p dùng đúng k và bốn notebook có bản executed.", first_line=False)

    # TÀI LIỆU THAM KHẢO
    add_heading(doc, "TÀI LIỆU THAM KHẢO", 1, page_break=True)
    references = [
        "[1] MLG-ULB, Credit Card Fraud Detection dataset. https://www.kaggle.com/datasets/mlg-ulb/creditcardfraud",
        "[2] TensorFlow, Classification on imbalanced data. https://www.tensorflow.org/tutorials/structured_data/imbalanced_data",
        "[3] scikit-learn, average_precision_score. https://scikit-learn.org/stable/modules/generated/sklearn.metrics.average_precision_score.html",
        "[4] scikit-learn, precision_recall_curve. https://scikit-learn.org/stable/modules/generated/sklearn.metrics.precision_recall_curve.html",
        "[5] scikit-learn, train_test_split. https://scikit-learn.org/stable/modules/generated/sklearn.model_selection.train_test_split.html",
        "[6] scikit-learn, Pipeline. https://scikit-learn.org/stable/modules/generated/sklearn.pipeline.Pipeline.html",
        "[7] scikit-learn, DummyClassifier. https://scikit-learn.org/stable/modules/generated/sklearn.dummy.DummyClassifier.html",
        "[8] scikit-learn, LogisticRegression. https://scikit-learn.org/stable/modules/generated/sklearn.linear_model.LogisticRegression.html",
        "[9] scikit-learn, RandomForestClassifier. https://scikit-learn.org/stable/modules/generated/sklearn.ensemble.RandomForestClassifier.html",
        "[10] Fraud Detection Handbook. https://fraud-detection-handbook.github.io/fraud-detection-handbook/Foreword.html",
        "[11] GitHub Docs, About large files on GitHub. https://docs.github.com/en/repositories/working-with-files/managing-large-files/about-large-files-on-github",
        "[12] GitHub Docs, Ignoring files. https://docs.github.com/en/get-started/git-basics/ignoring-files",
    ]
    for reference in references:
        paragraph = add_body(doc, reference, first_line=False)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        paragraph.paragraph_format.line_spacing = 1.0

    # Ensure styles remain explicit and fields refresh in Word.
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.text:
                format_run(run)
    set_update_fields(doc)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
